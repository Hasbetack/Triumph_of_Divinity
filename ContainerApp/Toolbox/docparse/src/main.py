from docx import Document
import json
import re
import os

rootdir = '\\'

document = Document(os.path.join(rootdir, "test.docx"))
tables = document.tables
datacards = {}
factions_durability = {}


def parse_stats(stats_table):
    stats = {}
    for col in stats_table.columns:
        col_cells = col.cells
        stats[col_cells[0].text] = col_cells[1].text
    return stats


def parse_abilities(abilities_table):
    row = abilities_table.rows[0]
    abilities_text = row.cells[0].text
    return abilities_text


def parse_caster(abilities_table):
    if len(abilities_table.columns) == 2:
        column = abilities_table.columns[1]
        caster_text = column.cells[0].text
        return caster_text
    return ""


def weapon_sanitizer(weapon):

    abilities = weapon["Abilities"]
    abilities_list = abilities.split(", ")
    weapon["Abilities"] = abilities_list

    return weapon


def sanitize_weapon_stats(weapon):
    weapon["Name"] = str(weapon["Name"]).replace("\n", "")
    weapon["Cost"] = int(weapon["Cost"])
    weapon["Range"] = str(weapon["Range"]).replace("\u201d", "")
    weapon["Wound"] = str(weapon["Wound"]).replace("+", "")
    if len(weapon["Rend"]) > 1:
        weapon["Rend"] = str(weapon["Rend"]).replace("-", "")
    else:
        weapon["Rend"] = ""

    return weapon


def weapon_maker(weapons_table, weapon_count_indices):
    weapon = {}
    weapon_name = weapons_table.rows[weapon_count_indices[0]].cells[0].text

    weapon_headers = []
    for cell in weapons_table.rows[weapon_count_indices[1]].cells:
        weapon_headers.append(cell.text)

    weapon_stats = []
    for cell in weapons_table.rows[weapon_count_indices[2]].cells:
        weapon_stats.append(cell.text)

    weapon["Name"] = weapon_name
    for index, header in enumerate(weapon_headers):
        weapon[header] = weapon_stats[index]

    weapon = weapon_sanitizer(weapon)
    weapon = sanitize_weapon_stats(weapon)

    return weapon


def parse_weapons(weapons_table):
    weapons = []
    row_index = len(weapons_table.rows)

    weapons.append(weapon_maker(weapons_table, [0, 1, 2]))

    if row_index == 6:
        weapons.append(weapon_maker(weapons_table, [3, 4, 5]))

    if row_index == 9:
        weapons.append(weapon_maker(weapons_table, [6, 7, 8]))

    return weapons


def parse_header(header):
    lines = header.split('\n')
    text_lines = list(filter(None, lines))

    first_line = text_lines[0].split(r'/')
    name = ' '.join(first_line[0].split()[:-2])
    points = first_line[0].split()[-2]

    unit_size = first_line[1].split()[0]
    if '-' in unit_size:
        unit_size_min, unit_size_max = unit_size.split('-')
    else:
        unit_size_min = unit_size
        unit_size_max = unit_size

    base_size = first_line[2].split()[0]

    keywords = text_lines[1].split(',')
    keywords = [keyword.strip() for keyword in keywords]
    faction = keywords[0]

    unit_loadout = text_lines[2]

    return name, points, faction, unit_size_min, unit_size_max, base_size, keywords, unit_loadout


def parse_datacard(datacard):
    row = datacard.rows[0]
    cell = row.cells[0]
    name, points, faction, unit_size_min, unit_size_max, base_size, keywords, unit_loadout = parse_header(cell.text)

    stats_table, weapons_table, abilities_table = cell.tables
    stats = parse_stats(stats_table)
    weapons = parse_weapons(weapons_table)
    abilities = parse_abilities(abilities_table)
    caster = parse_caster(abilities_table)
    return name, faction, points, unit_size_min, unit_size_max, base_size, keywords, unit_loadout, stats, weapons, abilities, caster


def statline_sanitizer(stats):

    sanitized_movement = stats["Move"].replace("”", "")
    stats["Move"] = int(sanitized_movement)

    sanitized_dash = stats["Dash"].replace("”", "")
    if sanitized_dash == "-":
        stats["Dash"] = -1
    else:
        stats["Dash"] = int(sanitized_dash)

    sanitized_melee = stats["Melee"].replace("+", "")
    if sanitized_melee == "-":
        stats["Melee"] = -1
    else:
        stats["Melee"] = int(sanitized_melee)

    sanitized_ranged = stats["Ranged"].replace("+", "")
    if sanitized_ranged == "-":
        stats["Ranged"] = -1
    else:
        stats["Ranged"] = int(sanitized_ranged)

    sanitized_durability = stats["Durability"].replace("+", "")
    if sanitized_durability == "-":
        stats["Durability"] = -1
    else:
        stats["Durability"] = int(sanitized_durability)

    sanitized_attacks = stats["Attacks"]
    if "D" not in sanitized_attacks:
        stats["Attacks"] = int(sanitized_attacks)

    sanitized_armour = stats["Armour"].replace("+", "")
    if sanitized_armour == "-":
        stats["Armour"] = -1
    else:
        stats["Armour"] = int(sanitized_armour)

    return stats


def sanitize_abilities(abilities):

    abilities_list = abilities.split("\n")
    unit_abilities = []
    faction_abilities = []
    abilities_sanitized = []

    del abilities_list[0]

    for i, blank in enumerate(abilities_list):
        if blank == "":
            del abilities_list[i]

    for ability in abilities_list:
        if ":" in ability:
            ability_split = ability.split(":")
            unit_abilities.append({
                "Ability_Name": ability_split[0],
                "Ability_Effect": ability_split[1].replace("\u201d", "\""),
            })
        else:
            faction_abilities.append(ability)

    abilities_sanitized.append(unit_abilities)
    abilities_sanitized.append(faction_abilities)

    return abilities_sanitized


def sanitize_caster(caster):

    caster_list = caster.split("\n")
    spells = []
    spells_sanitized = []

    if caster_list[0] != '':
        del caster_list[0]
        expression1 = re.compile('(\s*)This Formation is a Level (\s*)')
        expression1a = re.compile('(\s*)This model is a Level (\s*)')
        expression2 = re.compile('(\s*) Caster(\s*)')

        level = expression1.sub('', caster_list[0])
        level = expression1a.sub('', level)
        level = expression2.sub('', level)

        for i, blank in enumerate(caster_list):
            if blank == '':
                del caster_list[i]

        for spell in caster_list:
            if ":" in spell:
                spell_split = spell.split(":")
                spells.append({
                    "Name": spell_split[0],
                    "Effect": spell_split[1].replace("\u201d", "\"")
                })

        spells_sanitized.append(level)
        spells_sanitized.append(spells)

        return spells_sanitized
    else:
        return [0, {}]


def make_dir(faction):
    if not os.path.isdir(faction):
        os.mkdir(faction)


def make_datacards():
    for table in tables[-85:]:
        name, faction, points, unit_size_min, unit_size_max, base_size, keywords, unit_loadout, stats, weapons, abilities, caster = parse_datacard(table)
        if faction not in datacards:
            datacards[faction] = {}

        make_dir(faction)

        stats = statline_sanitizer(stats)
        abilities = sanitize_abilities(abilities)
        caster = sanitize_caster(caster)
        datacard_key = name.replace(" ", "_").lower()

        datacards[faction][datacard_key] = {
            "Name": name,
            "Keywords": keywords,
            "Points": int(points),
            "Unit_Size": {
                "min": int(unit_size_min),
                "max": int(unit_size_max)
            },
            "Move": stats["Move"],
            "Dash": stats["Dash"],
            "Melee": stats["Melee"],
            "Ranged": stats["Ranged"],
            "Strength": int(stats["Strength"]),
            "Durability": stats["Durability"],
            "Health": int(stats["Health"]),
            "Attacks": stats["Attacks"],
            "Will": int(stats["Will"]),
            "Armour": stats["Armour"],
            "Equipment_Options": unit_loadout,
            "Weapons": weapons,
            "Faction_Abilities": abilities[1],
            "Unit_Abilities": abilities[0],
            "Caster": caster[0],
            "Spells": caster[1]
        }


make_datacards()
for datacard_faction_key, faction in datacards.items():
    for datacard_name_key, name in faction.items():
        with open(datacard_faction_key + "\\" + datacard_name_key + ".json", "w+") as datacard_json:
            datacard_json.write(json.dumps(name))
